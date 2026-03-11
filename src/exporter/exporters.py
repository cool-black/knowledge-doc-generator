"""
文档导出器
支持 Markdown、PDF、Word 格式
"""

import os
from abc import ABC, abstractmethod
from datetime import datetime
from typing import Optional

import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class BaseExporter(ABC):
    """导出器基类"""

    def __init__(self, config: dict):
        self.config = config

    @abstractmethod
    def export(self, content: str, output_path: str) -> str:
        """
        导出文档
        返回输出文件路径
        """
        pass


class MarkdownExporter(BaseExporter):
    """Markdown 导出器"""

    def export(self, content: str, output_path: str) -> str:
        """保存 Markdown 文件"""
        # 确保目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)

        return output_path


class PDFExporter(BaseExporter):
    """PDF 导出器"""

    def __init__(self, config: dict):
        super().__init__(config)
        self.weasyprint_available = None  # 延迟检查

    def _check_weasyprint(self):
        """检查 weasyprint 是否可用"""
        if self.weasyprint_available is None:
            try:
                from weasyprint import HTML, CSS
                self.weasyprint_available = True
            except Exception:
                self.weasyprint_available = False
        return self.weasyprint_available

    def export(self, content: str, output_path: str) -> str:
        """将 Markdown 转换为 PDF"""
        if not self._check_weasyprint():
            raise ImportError(
                "weasyprint 不可用。Windows 用户需要安装 GTK+，"
                "或使用 Markdown/Word 格式导出。"
                "详见: https://doc.courtbouillon.org/weasyprint/stable/first_steps.html#windows"
            )

        from weasyprint import HTML, CSS

        # Markdown 转 HTML
        html_content = markdown.markdown(
            content,
            extensions=['tables', 'fenced_code', 'toc']
        )

        # 添加样式
        full_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{
                    font-family: "Noto Sans CJK SC", "WenQuanYi Micro Hei", sans-serif;
                    font-size: 11pt;
                    line-height: 1.6;
                    max-width: 210mm;
                    margin: 0 auto;
                    padding: 20mm;
                }}
                h1 {{ font-size: 24pt; color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }}
                h2 {{ font-size: 18pt; color: #444; border-bottom: 1px solid #ccc; padding-bottom: 5px; margin-top: 30px; }}
                h3 {{ font-size: 14pt; color: #555; margin-top: 20px; }}
                code {{
                    background-color: #f4f4f4;
                    padding: 2px 5px;
                    border-radius: 3px;
                    font-family: "Consolas", "Monaco", monospace;
                }}
                pre {{
                    background-color: #f4f4f4;
                    padding: 15px;
                    border-radius: 5px;
                    overflow-x: auto;
                }}
                blockquote {{
                    border-left: 4px solid #ccc;
                    margin: 0;
                    padding-left: 20px;
                    color: #666;
                }}
                table {{
                    border-collapse: collapse;
                    width: 100%;
                    margin: 20px 0;
                }}
                th, td {{
                    border: 1px solid #ddd;
                    padding: 8px;
                    text-align: left;
                }}
                th {{
                    background-color: #f2f2f2;
                }}
            </style>
        </head>
        <body>
            {html_content}
        </body>
        </html>
        """

        os.makedirs(os.path.dirname(output_path), exist_ok=True)

        HTML(string=full_html).write_pdf(output_path)

        return output_path


class DocxExporter(BaseExporter):
    """Word 导出器"""

    def export(self, content: str, output_path: str) -> str:
        """将 Markdown 转换为 Word"""
        doc = Document()

        # 设置默认字体
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(11)

        # 简单解析 Markdown（支持标题、段落、代码块）
        lines = content.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 标题
            if line.startswith('# '):
                p = doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                p = doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                p = doc.add_heading(line[4:], level=3)
            # 代码块
            elif line.startswith('```'):
                code_lines = []
                i += 1
                while i < len(lines) and not lines[i].startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                code = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.shading.background_pattern = 1
            # 普通段落
            elif line.strip():
                # 处理粗体和斜体
                p = doc.add_paragraph(line)
            else:
                # 空行
                pass

            i += 1

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        return output_path


class DocumentExporter:
    """导出器统一接口"""

    FORMATS = {
        'markdown': MarkdownExporter,
        'pdf': PDFExporter,
        'docx': DocxExporter,
    }

    def __init__(self, config: dict):
        self.config = config
        self.exporters = {}

        for fmt, ExporterClass in self.FORMATS.items():
            self.exporters[fmt] = ExporterClass(config.get(fmt, {}))

    def export(
        self,
        content: str,
        output_path: Optional[str] = None,
        fmt: Optional[str] = None,
    ) -> str:
        """
        导出文档

        Args:
            content: Markdown 内容
            output_path: 输出路径（可选）
            fmt: 格式（可选，默认从 output_path 推断或配置）

        Returns:
            输出文件路径
        """
        # 确定格式
        if fmt is None:
            if output_path:
                ext = os.path.splitext(output_path)[1].lower()
                fmt_map = {'.md': 'markdown', '.pdf': 'pdf', '.docx': 'docx'}
                fmt = fmt_map.get(ext, self.config.get('default_format', 'markdown'))
            else:
                fmt = self.config.get('default_format', 'markdown')

        if fmt not in self.exporters:
            raise ValueError(f"不支持的格式: {fmt}")

        # 确定输出路径
        if output_path is None:
            output_dir = self.config.get('output_dir', './output')
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"document_{timestamp}.{fmt if fmt != 'markdown' else 'md'}"
            output_path = os.path.join(output_dir, filename)

        # 执行导出
        exporter = self.exporters[fmt]
        return exporter.export(content, output_path)
